from PyDictionary import PyDictionary
from docx import Document
import nltk
import os
import random

nltk.download('punkt')

files = os.listdir('texts')
fulltext = []
document = Document('texts/Text 1.docx')

for file in files:
    if(file.split('.')[1] == 'docx' or file.split('.')[1] == 'DOCX'):
        document = Document('texts/' + file)
        for para in document.paragraphs:
            text = para.text
            tokens = nltk.word_tokenize(text)
            fulltext = fulltext + tokens

unduplicated_list = set(fulltext)
unduplicated_list = list(unduplicated_list)
fulltext = []

cnt = 70 # reik 70 zodziu

wordlist = []
document = Document("anglu.docx")
table = document.add_table(rows=2, cols=2)
length = len(unduplicated_list)
while cnt != 0:
    rng = random.randint(0,length)
    word = unduplicated_list[rng]
    if(len(word) < 6):
        continue
    try:
        meaning = PyDictionary.meaning(word)
        wordlist.append(word)
        cells = table.add_row().cells
        cells[0].text = wordlist
        cells[1].text = meaning
        cnt = cnt - 1
    except:
        continue

